import logging
from pathlib import Path
from typing import List, Dict, Any
import pypdf
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document loading and preprocessing."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def load_document(self, file_path: Path) -> str:
        """Load document content based on file extension."""
        try:
            suffix = file_path.suffix.lower()
            
            if suffix == '.pdf':
                return self._load_pdf(file_path)
            elif suffix == '.txt':
                return self._load_txt(file_path)
            elif suffix in ['.docx', '.doc']:
                return self._load_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {suffix}")
                
        except Exception as e:
            logger.error(f"Error loading {file_path}: {str(e)}")
            raise
    
    def _load_pdf(self, file_path: Path) -> str:
        """Load PDF file."""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def _load_txt(self, file_path: Path) -> str:
        """Load text file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _load_docx(self, file_path: Path) -> str:
        """Load DOCX file."""
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def process_documents(self, document_paths: List[Path]) -> List[Document]:
        """Process multiple documents into chunks."""
        all_documents = []
        
        for path in document_paths:
            try:
                logger.info(f"Processing: {path.name}")
                content = self.load_document(path)
                
                # Create document with metadata
                doc = Document(
                    page_content=content,
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "file_type": path.suffix
                    }
                )
                
                # Split into chunks
                chunks = self.text_splitter.split_documents([doc])
                
                # Add chunk metadata
                for i, chunk in enumerate(chunks):
                    chunk.metadata.update({
                        "chunk_id": i,
                        "total_chunks": len(chunks)
                    })
                
                all_documents.extend(chunks)
                logger.info(f"Created {len(chunks)} chunks from {path.name}")
                
            except Exception as e:
                logger.error(f"Failed to process {path}: {str(e)}")
        
        return all_documents